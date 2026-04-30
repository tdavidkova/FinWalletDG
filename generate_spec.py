from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Title Page ---
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FinWallet DG')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('System Requirements Specification')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Web-Based Kindergarten Class Fund Management System\n').font.size = Pt(13)
meta.add_run(f'\nVersion 1.0\n{datetime.date.today().strftime("%B %d, %Y")}').font.size = Pt(11)

doc.add_page_break()

# --- Table of Contents placeholder ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    ('1', 'Introduction', '3'),
    ('2', 'System Overview', '3'),
    ('3', 'User Roles & Authentication', '4'),
    ('4', 'Functional Requirements', '5'),
    ('4.1', 'Student Management', '5'),
    ('4.2', 'Transaction Management (Deposits & Expenses)', '6'),
    ('4.3', 'Balance Overview', '7'),
    ('4.4', 'Expense Categories & Customization', '8'),
    ('4.5', 'Receipts & Invoices Registry', '9'),
    ('4.6', 'Multi-Currency Support', '9'),
    ('4.7', 'Student Lifecycle (Enrollment / Unenrollment)', '10'),
    ('4.8', 'Sibling Handling', '10'),
    ('4.9', 'Reporting & Export', '11'),
    ('5', 'Non-Functional Requirements', '11'),
    ('6', 'Data Model', '12'),
    ('7', 'UI Wireframe Descriptions', '14'),
    ('8', 'Glossary', '16'),
]
for num, title_text, page in toc_items:
    p = doc.add_paragraph()
    indent = Cm(1) if '.' in num else Cm(0)
    p.paragraph_format.left_indent = indent
    p.add_run(f'{num}  {title_text}')

doc.add_page_break()

# ===== 1. INTRODUCTION =====
doc.add_heading('1. Introduction', level=1)

doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'This document specifies the requirements for FinWallet DG — a web-based application '
    'that replaces the existing Excel-based kindergarten class fund tracker '
    '("ДГ Мечо Пух група 2А касичка.xlsx"). The system will provide parents and class '
    'treasurers with a transparent, easy-to-use platform for managing pooled class funds: '
    'tracking deposits from parents, recording shared expenses, maintaining per-student '
    'balances, and archiving receipts/invoices.'
)

doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph(
    'The application covers the full lifecycle of class fund management for one or more '
    'kindergarten groups. It supports multi-currency operations (BGN and EUR), per-child '
    'cost splitting, student enrollment/unenrollment workflows, sibling linking, and '
    'invoice tracking. It is designed to be used primarily on mobile devices but must '
    'also work on desktop browsers.'
)

doc.add_heading('1.3 Definitions & Abbreviations', level=2)
table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
terms = [
    ('Касичка', 'Class piggy bank — the pooled fund for a group'),
    ('Захранване', 'Top-up / deposit made by a parent into the fund'),
    ('Движение', 'Transaction — any deposit or expense entry'),
    ('Наличност', 'Balance — the current available amount per student'),
    ('Отписано дете', 'Unenrolled child — a student who has left the group'),
    ('Индивидуален разход', 'Per-child cost — a group expense divided among students'),
]
for i, (term, defn) in enumerate(terms):
    table.rows[i].cells[0].text = term
    table.rows[i].cells[1].text = defn

doc.add_page_break()

# ===== 2. SYSTEM OVERVIEW =====
doc.add_heading('2. System Overview', level=1)
doc.add_paragraph(
    'FinWallet DG is a multi-tenant web application where each kindergarten group '
    '(e.g., "2А", "1А") operates as an independent fund. A class treasurer creates '
    'and administers the fund, invites parents, records transactions, and generates reports.'
)

doc.add_heading('2.1 Key Capabilities (derived from the existing Excel tracker)', level=2)
capabilities = [
    'Record per-student deposits (top-ups) with date and amount in BGN and/or EUR.',
    'Record group expenses that are automatically split equally among participating students.',
    'Maintain a real-time per-student balance in both BGN and EUR.',
    'Track children who leave the group, calculating their remaining balance for refund.',
    'Link siblings so they share a combined numbered slot and balance.',
    'Maintain a registry of receipts and invoices for each expense, including invoice numbers, '
    'dates, descriptions, totals, per-child cost, and currency.',
    'Separate receipt registries per group (e.g., "2A" vs "1A").',
    'Show the total fund balance across all active students.',
    'Support a configurable EUR ⇔ BGN exchange rate.',
    'Allow custom expense categories (theater, gifts, furniture, supplies, etc.).',
]
for cap in capabilities:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_page_break()

# ===== 3. USER ROLES =====
doc.add_heading('3. User Roles & Authentication', level=1)

doc.add_heading('3.1 Roles', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Role'
table.rows[0].cells[1].text = 'Description'
table.rows[0].cells[2].text = 'Permissions'
roles = [
    ('Administrator', 'System-wide admin', 'Full access; manage groups, users, settings'),
    ('Treasurer', 'Class fund manager (typically a parent elected by the group)',
     'Create/manage group; add/remove students; record deposits & expenses; manage invoices; generate reports; configure exchange rate'),
    ('Parent (Viewer)', 'Parent of a student in the group',
     'View own child\'s balance and transaction history; view group expense summaries; download reports'),
]
for i, (role, desc, perms) in enumerate(roles):
    table.rows[i+1].cells[0].text = role
    table.rows[i+1].cells[1].text = desc
    table.rows[i+1].cells[2].text = perms

doc.add_heading('3.2 Authentication', level=2)
auth_items = [
    'Email/password registration with email verification.',
    'Optional social login (Google).',
    'Invitation-based onboarding: Treasurer sends invite link to parents.',
    'Session management with JWT tokens; automatic logout after inactivity.',
]
for item in auth_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 4. FUNCTIONAL REQUIREMENTS =====
doc.add_heading('4. Functional Requirements', level=1)

# 4.1 Student Management
doc.add_heading('4.1 Student Management', level=2)
doc.add_paragraph(
    'The system shall maintain a roster of students belonging to each group.'
)

table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-SM-01', 'Treasurer can add a new student with full name.', 'Must'),
    ('FR-SM-02', 'Treasurer can edit a student\'s name.', 'Must'),
    ('FR-SM-03', 'Treasurer can mark a student as "unenrolled" (отписано дете), recording the unenrollment date and destination group (if transferring).', 'Must'),
    ('FR-SM-04', 'Unenrolled students appear in a separate "Unenrolled Children" section with their remaining balance.', 'Must'),
    ('FR-SM-05', 'Treasurer can link two or more students as siblings. Linked siblings share a single numbered slot and display a combined balance.', 'Must'),
    ('FR-SM-06', 'Students are automatically numbered sequentially within each group.', 'Should'),
    ('FR-SM-07', 'Treasurer can re-enroll a previously unenrolled student.', 'Could'),
    ('FR-SM-08', 'System supports bulk import of students from a CSV or Excel file.', 'Could'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# 4.2 Transaction Management
doc.add_heading('4.2 Transaction Management (Deposits & Expenses)', level=2)
doc.add_paragraph(
    'Transactions are the core of the system. Each transaction is either a deposit (positive) '
    'or an expense (negative) associated with one or more students.'
)

doc.add_heading('4.2.1 Deposits', level=3)
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-TD-01', 'Treasurer can record a deposit for a specific student, entering: amount, currency (BGN or EUR), date, and reason (default: "захранване" / top-up).', 'Must'),
    ('FR-TD-02', 'When a deposit is entered in one currency, the system automatically calculates the equivalent in the other currency using the configured exchange rate.', 'Must'),
    ('FR-TD-03', 'Treasurer can record a batch deposit (same amount for multiple students at once).', 'Should'),
    ('FR-TD-04', 'Deposits are reflected immediately in the student\'s balance.', 'Must'),
    ('FR-TD-05', 'Treasurer can record a refund ("възстановяване на сума") as a special negative deposit to return money to a parent.', 'Must'),
    ('FR-TD-06', 'Each deposit entry is immutable after creation. Corrections are made by adding a compensating transaction.', 'Should'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_heading('4.2.2 Expenses', level=3)
table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-TE-01', 'Treasurer can record a group expense by entering: total amount, currency, date, description/reason, and selecting which students participate.', 'Must'),
    ('FR-TE-02', 'The system automatically calculates the per-child cost by dividing the total equally among selected students.', 'Must'),
    ('FR-TE-03', 'The per-child cost is deducted from each participating student\'s balance.', 'Must'),
    ('FR-TE-04', 'The expense description supports free-text entry and an optional category tag.', 'Must'),
    ('FR-TE-05', 'Treasurer can select "all active students" or pick individual students for an expense.', 'Must'),
    ('FR-TE-06', 'The system supports linking an expense to one or more invoice/receipt records.', 'Should'),
    ('FR-TE-07', 'Treasurer can record an expense that applies to a subset of students (e.g., a theater trip that not all children attended).', 'Must'),
    ('FR-TE-08', 'System warns if an expense would cause any student\'s balance to go negative.', 'Should'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# 4.3 Balance Overview
doc.add_heading('4.3 Balance Overview', level=2)
doc.add_paragraph(
    'The balance view provides a snapshot of the current financial state of the fund.'
)
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-BO-01', 'Display a numbered list of all active students with their current balance in EUR (primary) and BGN (converted).', 'Must'),
    ('FR-BO-02', 'Display the total fund balance (sum of all active student balances) in both EUR and BGN.', 'Must'),
    ('FR-BO-03', 'Display a separate section for unenrolled children with their remaining balances.', 'Must'),
    ('FR-BO-04', 'For unenrolled children, show their original group and remaining balance in both currencies.', 'Must'),
    ('FR-BO-05', 'Balances update in real-time as transactions are added.', 'Must'),
    ('FR-BO-06', 'Clicking on a student\'s name navigates to their full transaction history.', 'Should'),
    ('FR-BO-07', 'Balance view supports sorting by name, balance amount, or student number.', 'Could'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# 4.4 Expense Categories
doc.add_heading('4.4 Expense Categories & Customization', level=2)
doc.add_paragraph(
    'Based on the existing tracker, expenses fall into recurring categories. '
    'The system shall support customizable categories.'
)
doc.add_heading('4.4.1 Default Categories (from existing data)', level=3)
categories = [
    ('Theater visits', 'Monthly theater outings — "Театър Октомври", "Театър Ноември", etc.'),
    ('Furniture & Equipment', 'Tables, chairs, bathroom mats — "11 маси и 2 столчета"'),
    ('Gifts for Staff', 'Christmas, end-of-year, and appreciation gifts — orchids, vouchers, candles, coffee'),
    ('Supplies & Materials', 'Craft supplies, glue, Christmas candles, work materials'),
    ('Graduation', 'End-of-year gifts and flowers for graduating classes'),
    ('Refund', 'Money returned to parents — "възстановяване на сума"'),
]
table = doc.add_table(rows=len(categories)+1, cols=2, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Category'
table.rows[0].cells[1].text = 'Examples from existing data'
for i, (cat, ex) in enumerate(categories):
    table.rows[i+1].cells[0].text = cat
    table.rows[i+1].cells[1].text = ex

doc.add_paragraph()
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-EC-01', 'Treasurer can create, edit, and delete custom expense categories.', 'Must'),
    ('FR-EC-02', 'Each expense can optionally be tagged with one category.', 'Should'),
    ('FR-EC-03', 'System provides the default categories listed above on group creation.', 'Should'),
    ('FR-EC-04', 'Expenses can be filtered by category in the transaction list and reports.', 'Should'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# 4.5 Receipts & Invoices
doc.add_heading('4.5 Receipts & Invoices Registry', level=2)
doc.add_paragraph(
    'A dedicated registry tracks all receipts and invoices associated with expenses, '
    'providing an audit trail. This mirrors the "Касови бележки и фактури" sheets.'
)
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-RI-01', 'Each receipt/invoice record contains: total amount, per-child cost, description, date, invoice number(s), and currency.', 'Must'),
    ('FR-RI-02', 'Receipts/invoices are linked to the corresponding expense transaction(s).', 'Must'),
    ('FR-RI-03', 'Treasurer can upload a photo or PDF scan of the physical receipt/invoice.', 'Should'),
    ('FR-RI-04', 'The registry is filterable by date range, category, and currency.', 'Should'),
    ('FR-RI-05', 'Each group has its own independent invoice registry.', 'Must'),
    ('FR-RI-06', 'The registry displays a summary row showing cumulative totals.', 'Should'),
    ('FR-RI-07', 'Invoice numbers are searchable.', 'Could'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

# 4.6 Multi-Currency
doc.add_heading('4.6 Multi-Currency Support', level=2)
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-MC-01', 'The system supports two currencies: BGN and EUR.', 'Must'),
    ('FR-MC-02', 'Treasurer can set and update the EUR ⇔ BGN exchange rate (default: 1.95583, the fixed rate).', 'Must'),
    ('FR-MC-03', 'All transactions and balances are stored and displayed in both currencies.', 'Must'),
    ('FR-MC-04', 'Currency conversion is applied automatically using the configured rate at time of entry.', 'Must'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# 4.7 Student Lifecycle
doc.add_heading('4.7 Student Lifecycle (Enrollment / Unenrollment)', level=2)
doc.add_paragraph(
    'Children may join or leave a group during the school year. The system must handle '
    'these transitions gracefully, preserving financial records and calculating refunds.'
)
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-SL-01', 'When a student is unenrolled, their balance freezes and they stop participating in new group expenses.', 'Must'),
    ('FR-SL-02', 'Unenrolled students are moved to a separate "Unenrolled Children" list showing: name, original group, remaining balance (BGN and EUR).', 'Must'),
    ('FR-SL-03', 'The total remaining balance of unenrolled children is displayed as a summary.', 'Must'),
    ('FR-SL-04', 'Treasurer can record deductions against the unenrolled pool (e.g., shared supplies purchased from leftover funds).', 'Should'),
    ('FR-SL-05', 'Treasurer can mark an unenrolled child\'s balance as "refunded" or "donated to fund".', 'Should'),
    ('FR-SL-06', 'When a child transfers to another group within the same system, their balance can be transferred automatically.', 'Could'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

# 4.8 Sibling Handling
doc.add_heading('4.8 Sibling Handling', level=2)
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-SH-01', 'Two or more students can be linked as siblings within a group.', 'Must'),
    ('FR-SH-02', 'Linked siblings share a single numbered position in the balance list.', 'Must'),
    ('FR-SH-03', 'A combined balance is shown for siblings, alongside individual breakdowns.', 'Should'),
    ('FR-SH-04', 'Expenses still apply individually per child; the sibling link is for display and reporting only.', 'Must'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# 4.9 Reporting & Export
doc.add_heading('4.9 Reporting & Export', level=2)
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('FR-RE-01', 'Export the full transaction ledger to Excel (.xlsx) format, matching the structure of the original "Движения" sheet.', 'Must'),
    ('FR-RE-02', 'Export the balance overview to Excel, matching the "Наличности" sheet.', 'Must'),
    ('FR-RE-03', 'Export the receipts/invoices registry to Excel, matching the "Касови бележки и фактури" sheets.', 'Must'),
    ('FR-RE-04', 'Generate a PDF summary report for a selected date range showing: all transactions, per-student balances, and expense breakdown by category.', 'Should'),
    ('FR-RE-05', 'Treasurer can generate a per-student statement showing all deposits and expenses for a specific child.', 'Should'),
    ('FR-RE-06', 'All exports include both BGN and EUR amounts.', 'Must'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# ===== 5. NON-FUNCTIONAL REQUIREMENTS =====
doc.add_heading('5. Non-Functional Requirements', level=1)

table = doc.add_table(rows=11, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'ID'
table.rows[0].cells[1].text = 'Requirement'
table.rows[0].cells[2].text = 'Priority'
reqs = [
    ('NFR-01', 'The application must be responsive and usable on mobile devices (min width: 360px).', 'Must'),
    ('NFR-02', 'The UI must support Bulgarian language. English is optional.', 'Must'),
    ('NFR-03', 'Page load time must be under 2 seconds on a 4G connection.', 'Should'),
    ('NFR-04', 'All financial data must be stored with at least 2 decimal places precision.', 'Must'),
    ('NFR-05', 'The system must support at least 50 groups with 30 students each.', 'Must'),
    ('NFR-06', 'All data must be backed up daily with point-in-time recovery.', 'Must'),
    ('NFR-07', 'Communication must be encrypted (HTTPS/TLS 1.2+).', 'Must'),
    ('NFR-08', 'Passwords must be hashed using bcrypt or Argon2.', 'Must'),
    ('NFR-09', 'The system must comply with GDPR for handling children\'s personal data (names).', 'Must'),
    ('NFR-10', 'Audit log: all create/update/delete operations must be logged with timestamp, user, and action.', 'Should'),
]
for i, (rid, req, pri) in enumerate(reqs):
    table.rows[i+1].cells[0].text = rid
    table.rows[i+1].cells[1].text = req
    table.rows[i+1].cells[2].text = pri

doc.add_page_break()

# ===== 6. DATA MODEL =====
doc.add_heading('6. Data Model', level=1)
doc.add_paragraph(
    'The following entity-relationship overview describes the core data model.'
)

# Group
doc.add_heading('6.1 Group', level=2)
table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Description'
fields = [
    ('id', 'UUID', 'Primary key'),
    ('name', 'String', 'Group name (e.g., "2А")'),
    ('kindergarten_name', 'String', 'Kindergarten name (e.g., "ДГ Мечо Пух")'),
    ('exchange_rate', 'Decimal(10,5)', 'EUR to BGN rate (default: 1.95583)'),
    ('created_at', 'Timestamp', 'Creation date'),
]
for i, (f, t, d) in enumerate(fields):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

# Student
doc.add_heading('6.2 Student', level=2)
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Description'
fields = [
    ('id', 'UUID', 'Primary key'),
    ('group_id', 'UUID (FK)', 'Reference to Group'),
    ('full_name', 'String', 'Student full name'),
    ('display_number', 'Integer', 'Sequential number in the group list'),
    ('status', 'Enum', '"active" or "unenrolled"'),
    ('unenrolled_at', 'Date (nullable)', 'Date when the child left the group'),
    ('sibling_group_id', 'UUID (nullable)', 'Links siblings together'),
]
for i, (f, t, d) in enumerate(fields):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

# Transaction
doc.add_heading('6.3 Transaction', level=2)
table = doc.add_table(rows=10, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Description'
fields = [
    ('id', 'UUID', 'Primary key'),
    ('student_id', 'UUID (FK)', 'Reference to Student'),
    ('amount_bgn', 'Decimal(10,2)', 'Amount in BGN (positive=deposit, negative=expense)'),
    ('amount_eur', 'Decimal(10,2)', 'Amount in EUR'),
    ('date', 'Date', 'Transaction date'),
    ('reason', 'String', 'Description / reason for the transaction'),
    ('category_id', 'UUID (FK, nullable)', 'Optional reference to ExpenseCategory'),
    ('expense_batch_id', 'UUID (nullable)', 'Groups transactions belonging to the same group expense'),
    ('created_by', 'UUID (FK)', 'User who created the transaction'),
]
for i, (f, t, d) in enumerate(fields):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

doc.add_page_break()

# ExpenseCategory
doc.add_heading('6.4 ExpenseCategory', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Description'
fields = [
    ('id', 'UUID', 'Primary key'),
    ('group_id', 'UUID (FK)', 'Reference to Group'),
    ('name', 'String', 'Category name (e.g., "Theater", "Gifts for Staff")'),
]
for i, (f, t, d) in enumerate(fields):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

# Invoice
doc.add_heading('6.5 Invoice / Receipt', level=2)
table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Field'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Description'
fields = [
    ('id', 'UUID', 'Primary key'),
    ('group_id', 'UUID (FK)', 'Reference to Group'),
    ('expense_batch_id', 'UUID (FK)', 'Links to the expense transaction batch'),
    ('total_amount', 'Decimal(10,2)', 'Total invoice amount'),
    ('per_child_cost', 'Decimal(10,2)', 'Calculated per-child share'),
    ('description', 'String', 'What was purchased'),
    ('invoice_number', 'String (nullable)', 'Invoice/receipt number(s)'),
    ('currency', 'Enum', '"BGN" or "EUR"'),
]
for i, (f, t, d) in enumerate(fields):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

doc.add_page_break()

# ===== 7. UI WIREFRAME DESCRIPTIONS =====
doc.add_heading('7. UI Wireframe Descriptions', level=1)
doc.add_paragraph(
    'The following describes the key screens of the application.'
)

screens = [
    ('7.1 Dashboard (Home)', [
        'Group selector dropdown (for treasurers managing multiple groups).',
        'Summary card: Total fund balance in EUR and BGN.',
        'Summary card: Number of active students.',
        'Summary card: Number of unenrolled students with pending refunds.',
        'Recent transactions list (last 10 entries).',
        'Quick-action buttons: "Add Deposit", "Add Expense".',
    ]),
    ('7.2 Student Roster', [
        'Numbered table of all active students with columns: #, Name, Balance (EUR), Balance (BGN).',
        'Sibling groups visually indicated (shared row number, indented sub-rows).',
        'Separate collapsible section: "Unenrolled Children" with remaining balances.',
        'Total balance row at the bottom.',
        '"Add Student" button; inline edit/unenroll actions per student.',
    ]),
    ('7.3 Transaction Ledger', [
        'Scrollable table with columns: Student Name, Amount (BGN), Amount (EUR), Date, Reason.',
        'Color coding: green for deposits, red for expenses.',
        'Filters: date range, student name, category, transaction type (deposit/expense).',
        'Search box for free-text search in reason/description.',
        'Pagination or infinite scroll.',
    ]),
    ('7.4 Add Deposit Form', [
        'Student selector (dropdown or multi-select for batch deposits).',
        'Amount input with currency toggle (BGN / EUR).',
        'Auto-calculated equivalent in the other currency.',
        'Date picker (defaults to today).',
        'Reason field (defaults to "захранване").',
        'Submit and "Submit & Add Another" buttons.',
    ]),
    ('7.5 Add Expense Form', [
        'Total amount input with currency toggle.',
        'Student multi-select with "Select All Active" shortcut.',
        'Auto-calculated per-child cost preview.',
        'Date picker.',
        'Description / reason text area.',
        'Category dropdown (optional).',
        'Invoice number input (optional).',
        'File upload for receipt photo/scan (optional).',
        'Warning banner if any student would go negative.',
    ]),
    ('7.6 Receipts & Invoices Registry', [
        'Table with one column per expense: Total, Per-child cost, Description, Date, Invoice #, Currency.',
        'Horizontal scroll for many entries (matching the Excel layout).',
        'Click on any entry to view/edit details and see linked receipt file.',
        '"Add Invoice" button.',
        'Filter by date range and currency.',
    ]),
    ('7.7 Settings', [
        'Group name and kindergarten name.',
        'EUR ⇔ BGN exchange rate configuration.',
        'Expense category management (CRUD).',
        'Member management: invite parents, assign roles.',
    ]),
]

for title, items in screens:
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ===== 8. GLOSSARY =====
doc.add_heading('8. Glossary', level=1)
table = doc.add_table(rows=11, cols=2, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Term'
table.rows[0].cells[1].text = 'Definition'
glossary = [
    ('ДГ', 'Детска Градина — Kindergarten'),
    ('Касичка', 'Piggy bank — the class pooled fund'),
    ('Захранване', 'Top-up / deposit — money paid in by a parent'),
    ('Движение', 'Transaction — any financial entry (deposit or expense)'),
    ('Наличност', 'Balance — available amount per student'),
    ('Отписано дете', 'Unenrolled child — student who has left the group'),
    ('Касова бележка', 'Cash receipt'),
    ('Фактура', 'Invoice'),
    ('Индивидуален разход', 'Per-child cost — expense total divided by number of participants'),
    ('Група', 'Group — a class of children (e.g., 1А, 2А)'),
]
for i, (term, defn) in enumerate(glossary):
    table.rows[i+1].cells[0].text = term
    table.rows[i+1].cells[1].text = defn

# --- Save ---
output_path = r'c:\Users\tdavidkova\OneDrive - DXC Production\Documents\Projects\GenAI\FinWalletDG\FinWalletDG_SRS.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
