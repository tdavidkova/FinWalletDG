"""Generate a DOCX deployment guide for hosting FinWalletDG on Render."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for i in range(1, 4):
    hs = doc.styles[f"Heading {i}"]
    hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

# --- Title ---
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("FinWallet DG")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Deployment Guide — Render.com + Neon PostgreSQL")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f"Generated: {datetime.date.today().isoformat()}").font.size = Pt(10)

doc.add_page_break()

# --- Table of Contents placeholder ---
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Prerequisites",
    "2. Prepare the Repository",
    "3. Create Required Files",
    "4. Set Up Neon PostgreSQL (Free)",
    "5. Push to GitHub",
    "6. Create a Render Web Service",
    "7. Configure Environment",
    "8. Deploy",
    "9. Verify the Deployment",
    "10. Custom Domain (Optional)",
    "11. Troubleshooting",
]
for item in toc_items:
    doc.add_paragraph(item, style="List Number")

doc.add_page_break()

# ---- CONTENT ----

# 1
doc.add_heading("1. Prerequisites", level=1)
doc.add_paragraph(
    "Before deploying, make sure you have the following:"
)
prereqs = [
    "A free account at https://render.com",
    "A free account at https://neon.tech",
    "A GitHub (or GitLab) account",
    "Git installed locally",
    "The FinWalletDG project source code",
]
for p in prereqs:
    doc.add_paragraph(p, style="List Bullet")

# 2
doc.add_heading("2. Prepare the Repository", level=1)
doc.add_paragraph(
    "Render deploys from a Git repository. If your project is not already in a "
    "Git repo, initialise one:"
)
doc.add_paragraph("git init\ngit add .\ngit commit -m \"Initial commit\"", style="No Spacing")
doc.add_paragraph()
doc.add_paragraph(
    "Make sure the following files are present in the root of the repository. "
    "The next section explains how to create any that are missing."
)

# 3
doc.add_heading("3. Create Required Files", level=1)

# 3.1 requirements.txt
doc.add_heading("3.1  requirements.txt", level=2)
doc.add_paragraph(
    "This file already exists in the project. Verify it contains at minimum:"
)
doc.add_paragraph(
    "fastapi==0.115.0\n"
    "uvicorn[standard]==0.30.0\n"
    "sqlalchemy==2.0.35\n"
    "pydantic==2.9.0\n"
    "python-multipart==0.0.9\n"
    "openpyxl==3.1.5\n"
    "psycopg2-binary==2.9.9\n"
    "gunicorn==22.0.0",
    style="No Spacing",
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Important: ")
run.bold = True
p.add_run(
    "Add gunicorn (production server) and psycopg2-binary (PostgreSQL driver) to the list."
)

# 3.2 render.yaml (optional)
doc.add_heading("3.2  render.yaml  (Infrastructure-as-Code, optional)", level=2)
doc.add_paragraph(
    "You can add a render.yaml file to the repo root to automate service creation. "
    "This is optional — you can also configure everything through the Render dashboard."
)
doc.add_paragraph(
    'services:\n'
    '  - type: web\n'
    '    name: finwallet-dg\n'
    '    runtime: python\n'
    '    plan: free\n'
    '    buildCommand: pip install -r requirements.txt\n'
    '    startCommand: gunicorn app.main:app -w 2 -k uvicorn.workers.UvicornWorker --bind 0.0.0.0:$PORT\n'
    '    envVars:\n'
    '      - key: PYTHON_VERSION\n'
    '        value: 3.12.0',
    style="No Spacing",
)

# 3.3 .gitignore
doc.add_heading("3.3  .gitignore", level=2)
doc.add_paragraph("Make sure these entries are in your .gitignore:")
doc.add_paragraph(
    "__pycache__/\n"
    "*.pyc\n"
    "finwallet.db\n"
    "backups/\n"
    ".env",
    style="No Spacing",
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Note: ")
run.bold = True
p.add_run(
    "The SQLite database file (finwallet.db) should NOT be committed. "
    "Render's free tier uses an ephemeral filesystem — the database will reset on every deploy. "
    "For production persistence, consider upgrading to a Render Disk or switching to PostgreSQL."
)

# 4 — Neon
doc.add_heading("4. Set Up Neon PostgreSQL (Free)", level=1)
doc.add_paragraph(
    "Neon provides a free-forever PostgreSQL database (512 MB). "
    "This replaces the local SQLite database so your data persists between Render deploys."
)

doc.add_heading("4.1  Create a Neon project", level=2)
neon_steps = [
    "Go to https://console.neon.tech and sign up (GitHub login works)",
    'Click \"New Project\"',
    "Name: FinWalletDG",
    "Region: AWS eu-central-1 (Frankfurt) — match your Render region",
    'Click \"Create Project\"',
]
for i, s in enumerate(neon_steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("4.2  Copy the connection string", level=2)
doc.add_paragraph(
    "After creating the project, Neon shows a connection string. Copy it — it looks like:"
)
doc.add_paragraph(
    "postgresql://neondb_owner:abc123@ep-cool-name-123456.eu-central-1.aws.neon.tech/neondb?sslmode=require",
    style="No Spacing",
)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Keep this string safe. ")
run.bold = True
p.add_run(
    "You will set it as the DATABASE_URL environment variable in Render. "
    "Never commit it to your repository."
)

doc.add_heading("4.3  How the app connects", level=2)
doc.add_paragraph(
    "The app reads the DATABASE_URL environment variable at startup. "
    "If set, it connects to PostgreSQL (Neon). "
    "If not set (local development), it falls back to the local SQLite file finwallet.db. "
    "No code changes are needed — this is already configured in app/database.py."
)

# 5
doc.add_heading("5. Push to GitHub", level=1)
doc.add_paragraph("Create a repository on GitHub and push your code:")
doc.add_paragraph(
    "git remote add origin https://github.com/YOUR_USERNAME/FinWalletDG.git\n"
    "git branch -M main\n"
    "git push -u origin main",
    style="No Spacing",
)

# 6
doc.add_heading("6. Create a Render Web Service", level=1)
steps = [
    "Log in to https://dashboard.render.com",
    'Click "New" → "Web Service"',
    'Select "Build and deploy from a Git repository" and click "Next"',
    "Connect your GitHub account if not already connected",
    "Find and select the FinWalletDG repository",
    "Configure the service (see settings below)",
    'Click "Create Web Service"',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("Service Settings", level=2)
table = doc.add_table(rows=8, cols=2, style="Light Grid Accent 1")
settings = [
    ("Name", "finwallet-dg"),
    ("Region", "Frankfurt (EU Central)  — or your preferred region"),
    ("Branch", "main"),
    ("Runtime", "Python"),
    ("Build Command", "pip install -r requirements.txt"),
    ("Start Command", "gunicorn app.main:app -w 2 -k uvicorn.workers.UvicornWorker --bind 0.0.0.0:$PORT"),
    ("Plan", "Free  (or Starter for persistent disk)"),
]
# Header row
table.cell(0, 0).text = "Setting"
table.cell(0, 1).text = "Value"
for cell in table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for i, (k, v) in enumerate(settings, 1):
    table.cell(i, 0).text = k
    table.cell(i, 1).text = v

# 7
doc.add_heading("7. Configure Environment", level=1)
doc.add_paragraph(
    "In the Render dashboard, go to your service → Environment and set:"
)
env_table = doc.add_table(rows=4, cols=3, style="Light Grid Accent 1")
env_table.cell(0, 0).text = "Key"
env_table.cell(0, 1).text = "Value"
env_table.cell(0, 2).text = "Notes"
for cell in env_table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
env_table.cell(1, 0).text = "DATABASE_URL"
env_table.cell(1, 1).text = "postgresql://... (your Neon connection string)"
env_table.cell(1, 2).text = "From step 4.2. Paste the full string."
env_table.cell(2, 0).text = "PYTHON_VERSION"
env_table.cell(2, 1).text = "3.12.0"
env_table.cell(2, 2).text = "Ensures the correct Python version"
env_table.cell(3, 0).text = "PORT"
env_table.cell(3, 1).text = "(set automatically by Render)"
env_table.cell(3, 2).text = "Do not override"

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Important: ")
run.bold = True
p.add_run(
    "The DATABASE_URL is the critical variable. Without it the app falls back to SQLite "
    "which is wiped on every Render deploy. With Neon your data is stored externally and persists forever."
)

# 8
doc.add_heading("8. Deploy", level=1)
doc.add_paragraph(
    "After creating the service, Render will automatically build and deploy your app. "
    "You can monitor the build logs in the Render dashboard."
)
doc.add_paragraph(
    "Every time you push to the main branch, Render will automatically redeploy. "
    "You can also trigger a manual deploy from the dashboard."
)
doc.add_paragraph("The build typically takes 1–3 minutes.")

# 9
doc.add_heading("9. Verify the Deployment", level=1)
doc.add_paragraph(
    "Once the deploy completes, Render provides a URL like:"
)
doc.add_paragraph("https://finwallet-dg.onrender.com", style="No Spacing")
doc.add_paragraph()
doc.add_paragraph("Open the URL in your browser — you should see the FinWallet DG dashboard.")
doc.add_paragraph()
doc.add_paragraph("To test the API, visit:")
doc.add_paragraph("https://finwallet-dg.onrender.com/docs", style="No Spacing")
doc.add_paragraph()
doc.add_paragraph("This shows the auto-generated FastAPI Swagger documentation.")

# 10
doc.add_heading("10. Custom Domain (Optional)", level=1)
steps = [
    "In the Render dashboard, go to your service → Settings → Custom Domains",
    'Click "Add Custom Domain"',
    "Enter your domain (e.g. wallet.example.com)",
    "Add the CNAME record shown by Render to your DNS provider",
    "Render automatically provisions an SSL certificate via Let's Encrypt",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

# 11
doc.add_heading("11. Troubleshooting", level=1)

doc.add_heading("App not starting", level=2)
doc.add_paragraph(
    "Check the Render logs for errors. Common issues:\n"
    "  • Missing gunicorn in requirements.txt\n"
    "  • Wrong start command — make sure it references app.main:app\n"
    "  • Python version mismatch — set PYTHON_VERSION env var"
)

doc.add_heading("Database resets on every deploy", level=2)
doc.add_paragraph(
    "If you forgot to set DATABASE_URL, the app uses local SQLite which is wiped on each deploy. "
    "Set DATABASE_URL to your Neon connection string (see step 4.2)."
)

doc.add_heading("Neon connection refused", level=2)
doc.add_paragraph(
    "Ensure the connection string includes ?sslmode=require at the end. "
    "Also verify the Neon project is not paused (free projects pause after 5 minutes of inactivity "
    "but wake up automatically on the next connection)."
)

doc.add_heading("Static files not loading", level=2)
doc.add_paragraph(
    "FastAPI serves static files through the app.mount() call. "
    "Make sure the static/ directory is committed to the repository."
)

doc.add_heading("Slow cold starts", level=2)
doc.add_paragraph(
    "Free tier services spin down after 15 minutes of inactivity. "
    "The first request after that takes ~30 seconds to spin up. "
    "Upgrade to the Starter plan to keep the service always running."
)

# --- Save ---
doc.save("FinWalletDG_Render_Deployment.docx")
print("Created: FinWalletDG_Render_Deployment.docx")
